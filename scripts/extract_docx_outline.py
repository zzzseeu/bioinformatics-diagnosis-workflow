#!/usr/bin/env python3
"""Extract a lightweight analysis outline from a DOCX file."""

from __future__ import annotations

import argparse
import json
import posixpath
import re
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET


NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}

METHOD_PATTERNS = [
    ("差异表达分析", ("DEG", "DEGs", "差异表达", "差异基因")),
    ("WGCNA", ("WGCNA", "共表达")),
    ("机器学习诊断模型", ("机器学习", "LASSO", "SVM", "随机森林", "特征基因")),
    ("GO/KEGG富集分析", ("GO", "KEGG", "富集")),
    ("GSEA", ("GSEA", "基因集富集")),
    ("PPI网络分析", ("PPI", "STRING", "蛋白互作")),
    ("ROC与外部验证", ("ROC", "AUC", "外部验证", "验证集")),
    ("列线图与临床模型评估", ("列线图", "Nomogram", "DCA", "校准曲线")),
    ("免疫浸润分析", ("免疫浸润", "CIBERSORT", "xCell", "ssGSEA")),
    ("单细胞分析", ("单细胞", "scRNA", "Seurat")),
]


class DocxParseError(RuntimeError):
    """Raised when a DOCX package cannot be parsed as expected."""


def _load_docx_module():
    try:
        import docx  # type: ignore
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to parse DOCX files. "
            "Please run this script in an environment that has a DOCX parser installed."
        ) from exc
    return docx


def _text_from_xml(node: ET.Element) -> str:
    return "".join(text.text or "" for text in node.findall(".//w:t", NS)).strip()


def _resolve_relationship_target(target: str, target_mode: str | None = None) -> str | None:
    if target_mode == "External":
        return None
    if re.match(r"^[a-z][a-z0-9+.-]*:", target, re.IGNORECASE):
        return None
    if target.lstrip("/").startswith("../"):
        return None
    normalized = posixpath.normpath(target.lstrip("/"))
    if not normalized.startswith("word/"):
        normalized = posixpath.normpath(posixpath.join("word", normalized))
    if normalized == ".." or normalized.startswith("../"):
        return None
    return normalized


def _image_media_paths(node: ET.Element, rels: dict[str, str]) -> list[str]:
    media_paths = []
    for blip in node.findall(".//a:blip", NS):
        rel_id = blip.attrib.get(f"{{{NS['r']}}}embed") or blip.attrib.get(f"{{{NS['r']}}}link")
        media_path = rels.get(rel_id or "")
        if media_path:
            media_paths.append(media_path)
    return media_paths


def _read_relationships(docx_path: Path) -> dict[str, str]:
    try:
        with zipfile.ZipFile(docx_path) as archive:
            try:
                raw = archive.read("word/_rels/document.xml.rels")
            except KeyError:
                return {}
    except zipfile.BadZipFile as exc:
        raise DocxParseError(f"Not a valid DOCX zip package: {docx_path}") from exc
    try:
        root = ET.fromstring(raw)
    except ET.ParseError as exc:
        raise DocxParseError(f"Cannot parse DOCX relationship XML: {docx_path}") from exc
    rels: dict[str, str] = {}
    for rel in root.findall("rel:Relationship", NS):
        rel_id = rel.attrib.get("Id")
        target = rel.attrib.get("Target")
        if not rel_id or not target:
            continue
        resolved = _resolve_relationship_target(target, rel.attrib.get("TargetMode"))
        if resolved:
            rels[rel_id] = resolved
    return rels


def _iter_document_events(docx_path: Path):
    rels = _read_relationships(docx_path)
    try:
        with zipfile.ZipFile(docx_path) as archive:
            try:
                document_xml = archive.read("word/document.xml")
            except KeyError as exc:
                raise DocxParseError(f"DOCX is missing word/document.xml: {docx_path}") from exc
    except zipfile.BadZipFile as exc:
        raise DocxParseError(f"Not a valid DOCX zip package: {docx_path}") from exc
    try:
        root = ET.fromstring(document_xml)
    except ET.ParseError as exc:
        raise DocxParseError(f"Cannot parse DOCX document XML: {docx_path}") from exc

    body = root.find("w:body", NS)
    if body is None:
        return

    for child in body:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text = _text_from_xml(child)
            image_paths = _image_media_paths(child, rels)
            if text or not image_paths:
                yield {"kind": "paragraph", "text": text}
            for media_path in image_paths:
                yield {"kind": "image", "media_path": media_path, "current_text": text}
        elif tag == "tbl":
            rows = []
            for row in child.findall(".//w:tr", NS):
                rows.append([_text_from_xml(cell) for cell in row.findall("./w:tc", NS)])
            yield {"kind": "table", "rows": rows}
            table_text = "\n".join("\t".join(row) for row in rows)
            for media_path in _image_media_paths(child, rels):
                yield {"kind": "image", "media_path": media_path, "current_text": table_text}


def _extract_paragraphs_and_tables(docx_path: Path):
    docx = _load_docx_module()
    try:
        document = docx.Document(str(docx_path))
    except Exception as exc:
        raise DocxParseError(f"Cannot open DOCX with available parser: {docx_path}") from exc

    paragraphs = []
    title_candidates = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        paragraphs.append({"text": text, "style": style_name})
        if style_name.lower().startswith(("heading", "title")):
            title_candidates.append(text)

    if not title_candidates and paragraphs:
        title_candidates.append(paragraphs[0]["text"])

    tables = []
    for index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append({"index": index, "rows": rows})

    return paragraphs, title_candidates, tables


def _copy_media(docx_path: Path, media_path: str, media_dir: Path, used_names: set[str]) -> Path:
    media_dir.mkdir(parents=True, exist_ok=True)
    name = Path(media_path).name
    if name in used_names:
        stem = Path(name).stem
        suffix = Path(name).suffix
        counter = 2
        while f"{stem}_{counter}{suffix}" in used_names:
            counter += 1
        name = f"{stem}_{counter}{suffix}"
    used_names.add(name)
    destination = media_dir / name
    with zipfile.ZipFile(docx_path) as archive:
        if media_path not in archive.namelist():
            raise DocxParseError(f"DOCX media target is missing: {media_path}")
        with archive.open(media_path) as source, destination.open("wb") as target:
            shutil.copyfileobj(source, target)
    return destination


def _infer_figure_type(context: str) -> str:
    context_lower = context.lower()
    if "火山图" in context or "volcano" in context_lower:
        return "volcano plot"
    if "热图" in context or "heatmap" in context_lower:
        return "heatmap"
    if "roc" in context_lower:
        return "ROC curve"
    if "箱线图" in context or "boxplot" in context_lower or "box plot" in context_lower:
        return "box plot"
    if "森林图" in context or "forest plot" in context_lower:
        return "forest plot"
    if "韦恩" in context or "venn" in context_lower:
        return "venn diagram"
    return "unknown"


def _detect_project_id(text: str) -> str | None:
    patterns = [
        r"(?:项目编号|项目ID|project\s*(?:id|no\.?|number))\s*[:：]\s*([^\s;；,，。】\n]+)",
        r"\b(project-[A-Za-z0-9_.-]+)\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1)
    return None


def _unique_matches(pattern: str, text: str, flags: int = 0) -> list[str]:
    seen = set()
    values = []
    for match in re.findall(pattern, text, flags):
        value = match if isinstance(match, str) else match[0]
        if value not in seen:
            seen.add(value)
            values.append(value)
    return values


def _detect_method_modules(text: str) -> list[dict[str, object]]:
    modules = []
    for name, keywords in METHOD_PATTERNS:
        hits = [keyword for keyword in keywords if keyword.lower() in text.lower()]
        if hits:
            modules.append({"name": name, "evidence": hits})
    return modules


def _extract_figures(docx_path: Path, out_dir: Path) -> list[dict[str, object]]:
    figures = []
    media_dir = out_dir / "media"
    previous_text = ""
    used_names: set[str] = set()

    for event in _iter_document_events(docx_path) or []:
        if event["kind"] == "paragraph":
            text = str(event["text"])
            if text:
                previous_text = text
        elif event["kind"] == "table":
            table_text = "\n".join("\t".join(row) for row in event["rows"])
            if table_text:
                previous_text = table_text
        elif event["kind"] == "image":
            media_path = str(event["media_path"])
            try:
                extracted_path = _copy_media(docx_path, media_path, media_dir, used_names)
            except DocxParseError:
                continue
            current_text = str(event.get("current_text") or "")
            context = "\n".join(part for part in (previous_text, current_text) if part)
            figures.append(
                {
                    "index": len(figures) + 1,
                    "source_media_path": media_path,
                    "extracted_path": str(extracted_path.resolve()),
                    "extracted_relpath": str(extracted_path.relative_to(out_dir)),
                    "previous_text": previous_text,
                    "current_text": current_text,
                    "figure_type": _infer_figure_type(context),
                    "is_reference_only": True,
                }
            )
    return figures


def _uncertainty_flags(result: dict[str, object]) -> list[str]:
    flags = []
    if not result.get("project_id"):
        flags.append("project_id_not_detected")
    if not result.get("title_candidates"):
        flags.append("title_not_detected")
    if not result.get("dataset_ids"):
        flags.append("dataset_ids_not_detected")
    if not result.get("method_modules"):
        flags.append("method_modules_not_detected")
    return flags


def _write_outputs(result: dict[str, object], out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "outline.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    lines = [
        "# DOCX Analysis Outline",
        "",
        f"- Source file: {result['source_file']}",
        f"- Project ID: {result.get('project_id') or 'not detected'}",
        "",
        "## Title Candidates",
    ]
    lines.extend(f"- {title}" for title in result.get("title_candidates", []))
    lines.extend(["", "## Dataset IDs"])
    lines.extend(f"- {dataset_id}" for dataset_id in result.get("dataset_ids", []))
    lines.extend(["", "## PMIDs"])
    lines.extend(f"- {pmid}" for pmid in result.get("pmids", []))
    lines.extend(["", "## Detected Modules"])
    for module in result.get("method_modules", []):
        evidence = ", ".join(module.get("evidence", []))
        lines.append(f"- {module['name']}: {evidence}")
    lines.extend(["", "## Figures"])
    for figure in result.get("figures", []):
        lines.append(f"- Figure {figure['index']}: {figure['figure_type']} ({figure['extracted_path']})")
    lines.extend(["", "## Tables"])
    for table in result.get("tables", []):
        lines.append(f"- Table {table['index']}: {len(table['rows'])} rows")
    lines.extend(["", "## Uncertainty Flags"])
    lines.extend(f"- {flag}" for flag in result.get("uncertainty_flags", []))
    (out_dir / "outline.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def extract_docx(docx_path, out_dir) -> dict:
    """Extract paragraphs, tables, media, and key analysis planning hints."""
    docx_path = Path(docx_path).resolve()
    out_dir = Path(out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    paragraphs, title_candidates, tables = _extract_paragraphs_and_tables(docx_path)
    paragraph_text = "\n".join(item["text"] for item in paragraphs)
    table_text = "\n".join("\t".join(row) for table in tables for row in table["rows"])
    all_text = "\n".join(part for part in (paragraph_text, table_text) if part)

    result: dict[str, object] = {
        "source_file": str(docx_path),
        "source_name": docx_path.name,
        "project_id": _detect_project_id(all_text),
        "title_candidates": title_candidates,
        "paragraphs": paragraphs,
        "dataset_ids": _unique_matches(r"GSE\d+", all_text, re.IGNORECASE),
        "pmids": _unique_matches(r"\bPMID\s*[:：]?\s*(\d{6,10})\b", all_text, re.IGNORECASE),
        "method_modules": _detect_method_modules(all_text),
        "tables": tables,
        "figures": _extract_figures(docx_path, out_dir),
        "image_rule": "Embedded Word images are reference figure types only, not real analysis results.",
    }
    result["uncertainty_flags"] = _uncertainty_flags(result)
    _write_outputs(result, out_dir)
    return result


def main(argv=None) -> int:
    parser = argparse.ArgumentParser(description="Extract an analysis outline from a DOCX file.")
    parser.add_argument("docx_path", help="Path to the DOCX file to parse.")
    parser.add_argument("--out-dir", default="docx-outline", help="Directory for outline.json, outline.md, and media.")
    args = parser.parse_args(argv)

    extract_docx(args.docx_path, args.out_dir)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
